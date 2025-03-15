import pdfplumber
import re
import os
from telegram import Update
from telegram.ext import Application, CommandHandler, MessageHandler, filters, CallbackContext
from docx import Document

# --- Telegram Token ---
TOKEN = "Your_telegram_token"

# --- Regex patterns ---
email_pattern = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"
phone_pattern = r"\+?\d[\d -]{8,14}\d"
linkedin_pattern = r"https?://(www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+"
github_pattern = r"https?://(www\.)?github\.com/[a-zA-Z0-9_-]+"

# --- Function to extract text from PDF ---
def extract_text_from_pdf(pdf_path):
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

# --- Function to extract structured information ---
def extract_info(text):
    email = re.search(email_pattern, text)
    phone = re.search(phone_pattern, text)
    linkedin = re.search(linkedin_pattern, text)
    github = re.search(github_pattern, text)

    sections = {
        "Education": [],
        "Skills": [],
        "Work Experience": [],
        "Projects": [],
        "Certificates": [],
        "Relevant Coursework": []
    }

    lines = text.split("\n")
    current_section = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Xác định tiêu đề của từng phần
        if "Education" in line:
            current_section = "Education"
        elif "Skills" in line:
            current_section = "Skills"
        elif "Work Experience" in line or "Experience" in line:
            current_section = "Work Experience"
        elif "Projects" in line:
            current_section = "Projects"
        elif "Certificates" in line:
            current_section = "Certificates"
        elif "Relevant Coursework" in line:
            current_section = "Relevant Coursework"
        elif current_section:
            sections[current_section].append(line)

    return {
        "Email": email.group(0) if email else "Không tìm thấy",
        "SĐT": phone.group(0) if phone else "Không tìm thấy",
        "LinkedIn": linkedin.group(0) if linkedin else "Không tìm thấy",
        "GitHub": github.group(0) if github else "Không tìm thấy",
        "Education": sections["Education"] or ["Không tìm thấy"],
        "Skills": sections["Skills"] or ["Không tìm thấy"],
        "Work Experience": sections["Work Experience"] or ["Không tìm thấy"],
        "Projects": sections["Projects"] or ["Không tìm thấy"],
        "Certificates": sections["Certificates"] or ["Không tìm thấy"],
        "Relevant Coursework": sections["Relevant Coursework"] or ["Không tìm thấy"]
    }

# --- Function to save data to DOCX with the same name as PDF ---
def save_to_docx(extracted_info, pdf_filename):
    # Chuyển đổi tên file từ .pdf sang .docx
    docx_filename = os.path.splitext(pdf_filename)[0] + ".docx"
    doc = Document()
    
    doc.add_heading("Thông tin trích xuất từ CV", level=1)

    doc.add_paragraph(f"📧 Email: {extracted_info['Email']}")
    doc.add_paragraph(f"📱 SĐT: {extracted_info['SĐT']}")
    doc.add_paragraph(f"🔗 LinkedIn: {extracted_info['LinkedIn']}")
    doc.add_paragraph(f"🐙 GitHub: {extracted_info['GitHub']}")

    sections = [
        ("🎓 Học vấn", extracted_info["Education"]),
        ("🛠 Kỹ năng", extracted_info["Skills"]),
        ("💼 Kinh nghiệm làm việc", extracted_info["Work Experience"]),
        ("📂 Dự án", extracted_info["Projects"]),
        ("📜 Chứng chỉ", extracted_info["Certificates"]),
        ("📖 Khóa học liên quan", extracted_info["Relevant Coursework"]),
    ]

    for title, content in sections:
        doc.add_heading(title, level=2)
        for item in content:
            doc.add_paragraph(f"- {item}")

    file_path = f"./{docx_filename}"
    doc.save(file_path)
    return file_path

# --- Start Command ---
async def start(update: Update, context: CallbackContext) -> None:
    await update.message.reply_text("👋 Chào bạn! Hãy gửi file PDF CV của bạn để tôi trích xuất thông tin.")

# --- Handle PDF files ---
async def handle_document(update: Update, context: CallbackContext) -> None:
    file = await context.bot.get_file(update.message.document.file_id)
    pdf_filename = update.message.document.file_name
    file_path = f"./{pdf_filename}"
    
    await file.download_to_drive(file_path)
    text = extract_text_from_pdf(file_path)
    extracted_info = extract_info(text)

    # Lưu vào file DOCX có cùng tên với PDF
    save_to_docx(extracted_info, pdf_filename)

    # Trả lời thông tin trích xuất trên Telegram
    response = (
        f"📄 **Thông tin trích xuất từ CV:**\n\n"
        f"📧 **Email:** {extracted_info['Email']}\n"
        f"📱 **SĐT:** {extracted_info['SĐT']}\n"
        f"🔗 **LinkedIn:** {extracted_info['LinkedIn']}\n"
        f"🐙 **GitHub:** {extracted_info['GitHub']}\n\n"
        f"🎓 **Học vấn:**\n" + "".join([f"- {edu}\n" for edu in extracted_info["Education"]]) + "\n"
        f"🛠 **Kỹ năng:**\n" + "".join([f"- {skill}\n" for skill in extracted_info["Skills"]]) + "\n"
        f"💼 **Kinh nghiệm làm việc:**\n" + "".join([f"- {exp}\n" for exp in extracted_info["Work Experience"]]) + "\n"
        f"📂 **Dự án:**\n" + "".join([f"- {proj}\n" for proj in extracted_info["Projects"]]) + "\n"
        f"📜 **Chứng chỉ:**\n" + "".join([f"- {cert}\n" for cert in extracted_info["Certificates"]]) + "\n"
        f"📖 **Khóa học liên quan:**\n" + "".join([f"- {course}\n" for course in extracted_info["Relevant Coursework"]]) + "\n"
    )

    # Chia nhỏ tin nhắn nếu quá dài
    messages = [response[i:i+4000] for i in range(0, len(response), 4000)]
    for msg in messages:
        await update.message.reply_text(msg, parse_mode="Markdown")

# --- Main Function to Start Bot ---
def main():
    app = Application.builder().token(TOKEN).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.Document.PDF, handle_document))

    print("🚀 Bot đang chạy...")
    app.run_polling()

if __name__ == "__main__":
    main()
